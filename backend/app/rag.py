import csv
import hashlib
import io
import json
import math
import re
import uuid
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import fitz
import httpx
from docx import Document
from openai import OpenAI
from qdrant_client import QdrantClient
from qdrant_client.http import models

from .config import RAG_DIR, settings
from .schemas import DocumentSummary, RetrievedChunk


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".json", ".csv"}
INDEX_VERSION = "rag-qdrant-hybrid-v1"
MANIFEST_PATH = RAG_DIR / "manifest.json"


@dataclass
class IndexedDocument:
    document_id: str
    file_name: str
    file_type: str
    content_type: str
    chunk_count: int
    uploaded_at: str
    sha256: str
    source_path: str
    indexing_version: str = INDEX_VERSION
    status: str = "indexed"
    embedding_model: str = ""
    embedding_fallback_used: bool = False


@dataclass
class IndexedChunk:
    document_id: str
    file_name: str
    file_type: str
    chunk_id: str
    point_id: str
    chunk_index: int
    content: str
    tokens: list[str]
    uploaded_at: str
    source_path: str
    indexing_version: str = INDEX_VERSION
    embedding_model: str = ""
    embedding_fallback_used: bool = False
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class RetrievalCandidate:
    chunk: IndexedChunk
    dense_score: float = 0.0
    lexical_score: float = 0.0
    fused_score: float = 0.0
    rerank_score: float | None = None
    strategy: str = "hybrid"


class EmbeddingService:
    def __init__(self) -> None:
        self.model = settings.rag_embedding_model
        self.dimensions = settings.rag_embedding_dimensions

    def embed_texts(self, texts: list[str]) -> tuple[list[list[float]], bool]:
        if not texts:
            return [], False
        try:
            if not settings.openai_api_key:
                raise RuntimeError("OPENAI_API_KEY is not configured for remote embeddings.")
            with httpx.Client(
                verify=settings.openai_verify_ssl,
                timeout=settings.openai_timeout_seconds,
                trust_env=False,
            ) as http_client:
                client = OpenAI(
                    api_key=settings.openai_api_key,
                    base_url=settings.openai_base_url,
                    http_client=http_client,
                    timeout=settings.openai_timeout_seconds,
                )
                response = client.embeddings.create(model=self.model, input=texts)
            embeddings = [item.embedding for item in response.data]
            self._validate(embeddings)
            return embeddings, False
        except Exception as exc:
            if not settings.rag_allow_local_embedding_fallback:
                raise RuntimeError(f"Embedding generation failed: {exc}") from exc
            return [stable_hash_embedding(text, self.dimensions) for text in texts], True

    def _validate(self, embeddings: list[list[float]]) -> None:
        if len({len(vector) for vector in embeddings}) != 1:
            raise RuntimeError("Embedding provider returned inconsistent vector dimensions.")
        dimension = len(embeddings[0])
        if dimension != self.dimensions:
            raise RuntimeError(
                f"Embedding dimension mismatch. Expected {self.dimensions}, got {dimension}."
            )


class SimpleReranker:
    def rerank(self, query: str, candidates: list[RetrievalCandidate]) -> list[RetrievalCandidate]:
        query_tokens = tokenize(query)
        query_phrase = normalize_for_matching(query)
        for candidate in candidates:
            text = normalize_for_matching(candidate.chunk.content)
            phrase_bonus = 0.18 if query_phrase and query_phrase in text else 0.0
            coverage = token_coverage(query_tokens, set(candidate.chunk.tokens))
            candidate.rerank_score = round(
                (0.55 * candidate.fused_score) + (0.35 * coverage) + phrase_bonus,
                6,
            )
        return sorted(
            candidates,
            key=lambda item: item.rerank_score if item.rerank_score is not None else item.fused_score,
            reverse=True,
        )


class PersistentHybridRAGStore:
    def __init__(self) -> None:
        self.documents: dict[str, IndexedDocument] = {}
        self.chunks: dict[str, IndexedChunk] = {}
        self.embedding_service = EmbeddingService()
        self.reranker = SimpleReranker()
        self.qdrant_available = False
        self.last_error: str | None = None
        self.client = self._init_qdrant()
        self._load_manifest()

    def _init_qdrant(self) -> QdrantClient | None:
        try:
            client = QdrantClient(path=settings.rag_vector_store_path)
            collections = client.get_collections().collections
            exists = any(collection.name == settings.rag_collection_name for collection in collections)
            if not exists:
                client.create_collection(
                    collection_name=settings.rag_collection_name,
                    vectors_config=models.VectorParams(
                        size=settings.rag_embedding_dimensions,
                        distance=models.Distance.COSINE,
                    ),
                )
            self.qdrant_available = True
            return client
        except Exception as exc:
            self.qdrant_available = False
            self.last_error = f"Qdrant initialization failed: {exc}"
            return None

    def _load_manifest(self) -> None:
        if not MANIFEST_PATH.exists():
            return
        try:
            payload = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
            self.documents = {
                item["document_id"]: IndexedDocument(**item)
                for item in payload.get("documents", [])
                if item.get("document_id")
            }
            self.chunks = {
                item["chunk_id"]: IndexedChunk(**item)
                for item in payload.get("chunks", [])
                if item.get("chunk_id")
            }
        except Exception as exc:
            self.documents = {}
            self.chunks = {}
            self.last_error = f"Manifest load failed: {exc}"

    def _persist_manifest(self) -> None:
        payload = {
            "indexing_version": INDEX_VERSION,
            "persisted_at": now_iso(),
            "vector_store": "qdrant_local",
            "collection_name": settings.rag_collection_name,
            "documents": [asdict(document) for document in self.documents.values()],
            "chunks": [asdict(chunk) for chunk in self.chunks.values()],
        }
        temp_path = MANIFEST_PATH.with_suffix(".tmp")
        temp_path.write_text(json.dumps(payload, indent=2, ensure_ascii=True), encoding="utf-8")
        temp_path.replace(MANIFEST_PATH)

    def list_documents(self) -> list[DocumentSummary]:
        return [
            DocumentSummary(
                document_id=document.document_id,
                file_name=document.file_name,
                content_type=document.content_type,
                chunk_count=document.chunk_count,
            )
            for document in sorted(self.documents.values(), key=lambda item: item.file_name.lower())
            if document.status == "indexed"
        ]

    def ingest(self, file_name: str, content_type: str, payload: bytes, source_path: str | None = None) -> tuple[str, int]:
        if not self.client or not self.qdrant_available:
            raise RuntimeError(self.last_error or "Qdrant vector store is not available.")

        extension = Path(file_name).suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type '{extension or 'unknown'}'.")

        file_hash = hashlib.sha256(payload).hexdigest()
        duplicate = self._find_duplicate(file_hash)
        if duplicate:
            return duplicate.document_id, duplicate.chunk_count

        text = extract_text(file_name=file_name, extension=extension, payload=payload)
        normalized_text = normalize_whitespace(text)
        if len(normalized_text) < 20:
            raise ValueError("The uploaded file did not contain enough readable text to index.")

        chunk_texts = chunk_text(normalized_text, settings.chunk_size, settings.chunk_overlap)
        if not chunk_texts:
            raise ValueError("The uploaded file could not be chunked into retrievable text.")

        document_id = str(uuid.uuid4())
        uploaded_at = now_iso()
        source_ref = source_path or ""
        embeddings: list[list[float]] = []
        fallback_used = False
        for batch in batched(chunk_texts, settings.rag_indexing_batch_size):
            batch_embeddings, batch_fallback = self.embedding_service.embed_texts(batch)
            embeddings.extend(batch_embeddings)
            fallback_used = fallback_used or batch_fallback

        chunks = [
            IndexedChunk(
                document_id=document_id,
                file_name=file_name,
                file_type=extension,
                chunk_id=f"{document_id}:{index}",
                point_id=str(uuid.uuid5(uuid.NAMESPACE_URL, f"{document_id}:{index}")),
                chunk_index=index,
                content=content,
                tokens=sorted(tokenize(content)),
                uploaded_at=uploaded_at,
                source_path=source_ref,
                embedding_model=settings.rag_embedding_model,
                embedding_fallback_used=fallback_used,
                metadata={
                    "sha256": file_hash,
                    "content_type": content_type or "application/octet-stream",
                    "char_length": len(content),
                },
            )
            for index, content in enumerate(chunk_texts)
        ]

        points = [
            models.PointStruct(
                id=chunk.point_id,
                vector=embedding,
                payload=chunk_payload(chunk),
            )
            for chunk, embedding in zip(chunks, embeddings, strict=True)
        ]
        self.client.upsert(collection_name=settings.rag_collection_name, points=points)

        document = IndexedDocument(
            document_id=document_id,
            file_name=file_name,
            file_type=extension,
            content_type=content_type or "application/octet-stream",
            chunk_count=len(chunks),
            uploaded_at=uploaded_at,
            sha256=file_hash,
            source_path=source_ref,
            embedding_model=settings.rag_embedding_model,
            embedding_fallback_used=fallback_used,
        )
        self.documents[document_id] = document
        self.chunks.update({chunk.chunk_id: chunk for chunk in chunks})
        self._persist_manifest()
        return document_id, len(chunks)

    def retrieve(self, query: str, top_k: int) -> list[RetrievedChunk]:
        diagnostics = self.retrieve_with_diagnostics(query, top_k)
        return diagnostics["results"]

    def retrieve_with_diagnostics(self, query: str, top_k: int) -> dict[str, Any]:
        normalized_query = query.strip()
        if not normalized_query:
            return self._empty_retrieval("hybrid", "The retrieval query was empty.")
        if not self.chunks:
            return self._empty_retrieval("hybrid", "No indexed RAG documents are available.")

        candidate_k = max(top_k, settings.rag_candidate_k)
        dense_candidates, dense_fallback = self._dense_candidates(normalized_query, candidate_k)
        lexical_candidates = self._lexical_candidates(normalized_query, candidate_k)
        fused = self._fuse_candidates(dense_candidates, lexical_candidates)

        candidates = list(fused.values())
        candidates.sort(key=lambda item: item.fused_score, reverse=True)
        candidates = dedupe_candidates(candidates)[:candidate_k]

        rerank_applied = False
        if settings.rag_enable_rerank and candidates:
            candidates = self.reranker.rerank(normalized_query, candidates[: settings.rag_rerank_top_n]) + candidates[settings.rag_rerank_top_n :]
            rerank_applied = True

        results = [
            to_retrieved_chunk(
                candidate,
                position=index,
                strategy="hybrid+dense+lexical+rerank" if rerank_applied else "hybrid+dense+lexical",
                dense_fallback=dense_fallback,
            )
            for index, candidate in enumerate(candidates[:top_k], start=1)
            if candidate.fused_score > 0 or (candidate.rerank_score or 0) > 0
        ]

        strength = classify_strength(results)
        return {
            "results": results,
            "diagnostics": {
                "retrieval_strategy": "hybrid",
                "dense_enabled": True,
                "lexical_enabled": settings.rag_enable_hybrid,
                "rerank_enabled": settings.rag_enable_rerank,
                "rerank_applied": rerank_applied,
                "candidate_k": candidate_k,
                "top_k": top_k,
                "dense_weight": settings.rag_dense_weight,
                "lexical_weight": settings.rag_lexical_weight,
                "embedding_model": settings.rag_embedding_model,
                "embedding_fallback_used": dense_fallback,
                "result_strength": strength,
                "qdrant_available": self.qdrant_available,
            },
        }

    def _dense_candidates(self, query: str, limit: int) -> tuple[dict[str, float], bool]:
        if not self.client or not self.qdrant_available:
            return {}, True
        try:
            vectors, fallback = self.embedding_service.embed_texts([query])
            search_results = self.client.search(
                collection_name=settings.rag_collection_name,
                query_vector=vectors[0],
                limit=limit,
                with_payload=True,
            )
            dense = {}
            for item in search_results:
                payload = item.payload or {}
                chunk_id = payload.get("chunk_id")
                if chunk_id:
                    dense[str(chunk_id)] = max(0.0, float(item.score))
            return dense, fallback
        except Exception as exc:
            self.last_error = f"Dense retrieval failed: {exc}"
            return {}, True

    def _lexical_candidates(self, query: str, limit: int) -> dict[str, float]:
        query_tokens = tokenize(query)
        if not query_tokens:
            return {}
        document_count = max(len(self.chunks), 1)
        document_frequency: Counter[str] = Counter()
        for chunk in self.chunks.values():
            document_frequency.update(set(chunk.tokens))

        scores: dict[str, float] = {}
        for chunk in self.chunks.values():
            chunk_tokens = Counter(chunk.tokens)
            length_norm = max(len(chunk.tokens), 1)
            score = 0.0
            for token in query_tokens:
                if token not in chunk_tokens:
                    continue
                idf = math.log((document_count + 1) / (document_frequency[token] + 0.5)) + 1
                score += (chunk_tokens[token] / length_norm) * idf
            if score > 0:
                scores[chunk.chunk_id] = score
        return dict(sorted(scores.items(), key=lambda item: item[1], reverse=True)[:limit])

    def _fuse_candidates(self, dense: dict[str, float], lexical: dict[str, float]) -> dict[str, RetrievalCandidate]:
        dense_norm = normalize_scores(dense)
        lexical_norm = normalize_scores(lexical)
        all_ids = set(dense_norm) | set(lexical_norm)
        fused: dict[str, RetrievalCandidate] = {}
        for chunk_id in all_ids:
            chunk = self.chunks.get(chunk_id)
            if not chunk:
                continue
            dense_score = dense_norm.get(chunk_id, 0.0)
            lexical_score = lexical_norm.get(chunk_id, 0.0)
            if not settings.rag_enable_hybrid:
                lexical_score = 0.0
            fused_score = (settings.rag_dense_weight * dense_score) + (settings.rag_lexical_weight * lexical_score)
            fused[chunk_id] = RetrievalCandidate(
                chunk=chunk,
                dense_score=round(dense_score, 6),
                lexical_score=round(lexical_score, 6),
                fused_score=round(fused_score, 6),
            )
        return fused

    def _find_duplicate(self, sha256: str) -> IndexedDocument | None:
        return next((document for document in self.documents.values() if document.sha256 == sha256), None)

    def _empty_retrieval(self, strategy: str, message: str) -> dict[str, Any]:
        return {
            "results": [],
            "diagnostics": {
                "retrieval_strategy": strategy,
                "rerank_applied": False,
                "result_strength": "none",
                "message": message,
                "qdrant_available": self.qdrant_available,
                "last_error": self.last_error,
            },
        }

    def stats(self) -> dict[str, Any]:
        return {
            "documents": len(self.documents),
            "chunks": len(self.chunks),
            "collection_name": settings.rag_collection_name,
            "vector_store_path": settings.rag_vector_store_path,
            "embedding_model": settings.rag_embedding_model,
            "indexing_version": INDEX_VERSION,
            "qdrant_available": self.qdrant_available,
            "last_error": self.last_error,
            "manifest_path": str(MANIFEST_PATH),
        }

    def health(self) -> dict[str, Any]:
        return {
            "status": "ok" if self.qdrant_available else "degraded",
            "qdrant_available": self.qdrant_available,
            "documents_loaded": len(self.documents),
            "chunks_loaded": len(self.chunks),
            "last_error": self.last_error,
        }

    def rebuild(self) -> dict[str, Any]:
        # Replays persisted chunk metadata into Qdrant. It does not reparse source files;
        # uploaded source files plus manifest remain the local source of truth for now.
        if not self.client:
            self.client = self._init_qdrant()
        if not self.client:
            raise RuntimeError(self.last_error or "Qdrant vector store is not available.")
        if self.client.collection_exists(settings.rag_collection_name):
            self.client.delete_collection(settings.rag_collection_name)
        self.client.create_collection(
            collection_name=settings.rag_collection_name,
            vectors_config=models.VectorParams(size=settings.rag_embedding_dimensions, distance=models.Distance.COSINE),
        )
        chunks = list(self.chunks.values())
        reembedded = 0
        fallback_used = False
        for batch in batched(chunks, settings.rag_indexing_batch_size):
            vectors, batch_fallback = self.embedding_service.embed_texts([chunk.content for chunk in batch])
            fallback_used = fallback_used or batch_fallback
            points = [
                models.PointStruct(id=chunk.point_id, vector=vector, payload=chunk_payload(chunk))
                for chunk, vector in zip(batch, vectors, strict=True)
            ]
            self.client.upsert(collection_name=settings.rag_collection_name, points=points)
            reembedded += len(points)
        return {"documents": len(self.documents), "chunks_reindexed": reembedded, "embedding_fallback_used": fallback_used}


rag_store = PersistentHybridRAGStore()


def extract_text(file_name: str, extension: str, payload: bytes) -> str:
    if extension in {".txt", ".md"}:
        return payload.decode("utf-8", errors="ignore")
    if extension == ".json":
        data = json.loads(payload.decode("utf-8", errors="ignore"))
        return json.dumps(data, indent=2, ensure_ascii=True)
    if extension == ".csv":
        text = payload.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(text))
        rows = [", ".join(cell.strip() for cell in row) for row in reader]
        return "\n".join(rows)
    if extension == ".pdf":
        with fitz.open(stream=payload, filetype="pdf") as document:
            return "\n".join(page.get_text("text") for page in document)
    if extension == ".docx":
        document = Document(io.BytesIO(payload))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    raise ValueError(f"Unsupported file: {file_name}")


def normalize_whitespace(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def tokenize(text: str) -> set[str]:
    return set(re.findall(r"[a-zA-Z0-9_]{2,}", text.lower()))


def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    if chunk_size <= 0:
        return []
    if chunk_overlap >= chunk_size:
        chunk_overlap = max(0, chunk_size // 4)

    chunks: list[str] = []
    start = 0
    length = len(text)
    while start < length:
        target_end = min(length, start + chunk_size)
        end = target_end
        if target_end < length:
            window = text[start:target_end]
            paragraph_break = window.rfind("\n\n")
            sentence_break = max(window.rfind(". "), window.rfind("? "), window.rfind("! "))
            whitespace_break = window.rfind(" ")
            best_break = max(paragraph_break, sentence_break, whitespace_break)
            if best_break > chunk_size * 0.55:
                end = start + best_break + 1
        chunk = text[start:end].strip()
        if chunk and len(chunk) > 10:
            chunks.append(chunk)
        if end >= length:
            break
        start = max(end - chunk_overlap, start + 1)
    return chunks


def combine_retrieval_results(results: Iterable[RetrievedChunk]) -> str:
    rendered = []
    for index, item in enumerate(results, start=1):
        metadata = item.source_metadata or {}
        strategy = item.retrieval_strategy or "hybrid"
        rendered.append(
            f"[Source {index}: {item.file_name} | chunk={item.chunk_index} | score={item.score} | strategy={strategy}]\n"
            f"{item.content}\n"
            f"Metadata: document_id={item.document_id}, file_type={item.file_type or metadata.get('file_type', 'unknown')}"
        )
    return "\n\n".join(rendered)


def chunk_payload(chunk: IndexedChunk) -> dict[str, Any]:
    return {
        "document_id": chunk.document_id,
        "chunk_id": chunk.chunk_id,
        "file_name": chunk.file_name,
        "file_type": chunk.file_type,
        "chunk_index": chunk.chunk_index,
        "uploaded_at": chunk.uploaded_at,
        "source_path": chunk.source_path,
        "indexing_version": chunk.indexing_version,
        "embedding_model": chunk.embedding_model,
        "embedding_fallback_used": chunk.embedding_fallback_used,
    }


def stable_hash_embedding(text: str, dimensions: int) -> list[float]:
    vector = [0.0] * dimensions
    tokens = tokenize(text)
    if not tokens:
        return vector
    for token in tokens:
        digest = hashlib.sha256(token.encode("utf-8")).digest()
        index = int.from_bytes(digest[:4], "big") % dimensions
        sign = 1.0 if digest[4] % 2 == 0 else -1.0
        vector[index] += sign
    norm = math.sqrt(sum(value * value for value in vector)) or 1.0
    return [value / norm for value in vector]


def normalize_scores(scores: dict[str, float]) -> dict[str, float]:
    if not scores:
        return {}
    max_score = max(scores.values()) or 1.0
    return {key: round(max(0.0, value) / max_score, 6) for key, value in scores.items()}


def token_coverage(query_tokens: set[str], chunk_tokens: set[str]) -> float:
    if not query_tokens:
        return 0.0
    return len(query_tokens & chunk_tokens) / len(query_tokens)


def normalize_for_matching(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()


def dedupe_candidates(candidates: list[RetrievalCandidate]) -> list[RetrievalCandidate]:
    seen: set[str] = set()
    deduped: list[RetrievalCandidate] = []
    for candidate in candidates:
        signature = hashlib.sha1(candidate.chunk.content[:500].lower().encode("utf-8")).hexdigest()
        if signature in seen:
            continue
        seen.add(signature)
        deduped.append(candidate)
    return deduped


def to_retrieved_chunk(candidate: RetrievalCandidate, position: int, strategy: str, dense_fallback: bool) -> RetrievedChunk:
    score = candidate.rerank_score if candidate.rerank_score is not None else candidate.fused_score
    chunk = candidate.chunk
    return RetrievedChunk(
        document_id=chunk.document_id,
        file_name=chunk.file_name,
        chunk_id=chunk.chunk_id,
        score=round(float(score), 4),
        content=chunk.content,
        file_type=chunk.file_type,
        chunk_index=chunk.chunk_index,
        dense_score=candidate.dense_score,
        lexical_score=candidate.lexical_score,
        fused_score=candidate.fused_score,
        rerank_score=candidate.rerank_score,
        ranking_position=position,
        retrieval_strategy=strategy,
        source_metadata={
            **chunk.metadata,
            "uploaded_at": chunk.uploaded_at,
            "source_path": chunk.source_path,
            "indexing_version": chunk.indexing_version,
            "embedding_model": chunk.embedding_model,
        },
        fallback_flags={
            "embedding_fallback_used": chunk.embedding_fallback_used or dense_fallback,
            "dense_available": candidate.dense_score > 0,
            "lexical_available": candidate.lexical_score > 0,
        },
    )


def classify_strength(results: list[RetrievedChunk]) -> str:
    if not results:
        return "none"
    best = max(result.score for result in results)
    if best >= 0.65:
        return "strong"
    if best >= 0.25:
        return "moderate"
    return "weak"


def batched[T](items: list[T], size: int) -> Iterable[list[T]]:
    batch_size = max(size, 1)
    for index in range(0, len(items), batch_size):
        yield items[index : index + batch_size]


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()
